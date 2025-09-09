from pathlib import Path
from docx import Document

from audit_lib.enrich.docx_sections import load_docx_sections


def test_load_docx_sections_missing_headings(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("No headings here")
    path = tmp_path / "no_heads.docx"
    doc.save(path)
    rq, sections = load_docx_sections(path)
    assert rq == ""
    assert sections == []


def test_load_docx_sections_basic(tmp_path: Path) -> None:
    doc = Document()
    doc.add_heading("Question", level=1)
    doc.add_paragraph("Intro")
    doc.add_heading("Methods", level=2)
    doc.add_paragraph("Study methods")
    path = tmp_path / "with_heads.docx"
    doc.save(path)
    rq, sections = load_docx_sections(path)
    assert rq == "Question"
    assert sections[1]["title"] == "Methods"
    assert sections[1]["body"] == "Study methods"
