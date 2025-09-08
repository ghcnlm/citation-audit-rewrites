# src/audit_lib/enrich.py
# -*- coding: utf-8 -*-
"""
Enrichment library:
- Parse review .docx headings (H1/H2/H3/H4) and build section bodies
- Match each claim to the most similar section (TF-IDF cosine, with penalties for generic sections)
- Robustly resolve source_pdf_path by normalizing citations and filenames:
    * Hyphens vs. underscores ("CLEAR-AA" -> "CLEAR_AA")
    * Remove "et al.", strip parentheses, commas, ampersands, etc.
    * Narrative vs. parenthetical formats ("Masilela (2022)" vs "(Chirau et al., 2022)")
    * Org names and multi-word surnames (e.g., "de Haan")
    * Multi-citations: pick the year from the sub-citation that names the row’s target author
    * Prefer the FIRST/lead author when multiple authors are joined with "&"
    * Penalize wrong-year candidates
No external deps besides python-docx.
"""

from __future__ import annotations
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import csv
import math
import re
import string

# ---------------------------
# Basic text utils
# ---------------------------

_PUNCT_TABLE_SP = str.maketrans({c: " " for c in string.punctuation})
_PUNCT_TABLE_DROP = str.maketrans("", "", string.punctuation)

_STOP = {
    "the","and","a","an","of","to","in","for","on","with","as","by","from","at","that","this","these","those",
    "is","are","was","were","be","been","being","it","its","their","there","which","or","not","but","if","than",
    "can","may","might","should","would","could","will","shall","do","does","did","done","such","into","over",
    "about","across","per","vs","via","within","between","among","both","also","more","most","much","many","some",
    "any","each","other","another","however","therefore","thus","so","because","while","where","when",
}

_HEADING_PENALTY_PATTERNS = [
    r"\bexecutive\s+summary\b",
    r"\babstract\b",
    r"\boverview\b",
    # r"\bintroduction\b",  # enable if Introduction dominates too often
]

def _norm_spaces(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip())

def _tokens(text: str) -> List[str]:
    if not text:
        return []
    text = text.lower().translate(_PUNCT_TABLE_SP)
    toks = [t for t in text.split() if len(t) > 1 and t not in _STOP and t.isascii()]
    return toks

def _tf(tokens: List[str]) -> Dict[str, float]:
    tf: Dict[str, float] = {}
    for t in tokens:
        tf[t] = tf.get(t, 0.0) + 1.0
    n = float(len(tokens)) or 1.0
    for k in list(tf.keys()):
        tf[k] /= n
    return tf

def _cosine(a: Dict[str, float], b: Dict[str, float]) -> float:
    if not a or not b:
        return 0.0
    dot = 0.0
    for k, v in a.items():
        if k in b:
            dot += v * b[k]
    na = math.sqrt(sum(v*v for v in a.values()))
    nb = math.sqrt(sum(v*v for v in b.values()))
    if na == 0.0 or nb == 0.0:
        return 0.0
    return dot / (na * nb)

# ---------------------------
# DOCX parsing (requires python-docx)
# ---------------------------

def _load_docx_sections(docx_path: Path) -> Tuple[str, List[Dict]]:
    """
    Returns:
      research_question (H1 title or ""),
      sections: list of dicts:
        {
          "level": int (1..4),
          "title": str,
          "body": str,   # text between this heading and the next heading
        }
    """
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx is required. pip install python-docx") from e

    if not docx_path.exists():
        return "", []

    doc = Document(str(docx_path))

    heads = []  # (idx, level, title)
    for i, p in enumerate(doc.paragraphs):
        style_name = (getattr(p.style, "name", "") or "").lower()
        if style_name.startswith("heading"):
            m = re.search(r"(\d+)", style_name)
            if not m:
                continue
            lvl = int(m.group(1))
            if 1 <= lvl <= 4:
                title = _norm_spaces(p.text)
                if title:
                    heads.append((i, lvl, title))

    if not heads:
        return "", []

    rq = ""
    for _, lvl, title in heads:
        if lvl == 1:
            rq = title
            break

    sections: List[Dict] = []
    for j, (idx, lvl, title) in enumerate(heads):
        nxt = heads[j + 1][0] if j + 1 < len(heads) else len(doc.paragraphs)
        chunks = []
        for k in range(idx + 1, nxt):
            t = _norm_spaces(doc.paragraphs[k].text)
            if t:
                chunks.append(t)
        body = " ".join(chunks)
        sections.append({"level": lvl, "title": title, "body": body})

    return rq, sections

# ---------------------------
# Section scoring
# ---------------------------

def _heading_penalty_factor(title: str) -> float:
    t = (title or "").lower()
    for pat in _HEADING_PENALTY_PATTERNS:
        if re.search(pat, t):
            return 0.60
    return 1.0

def _level_boost(level: int) -> float:
    return {4: 1.15, 3: 1.10, 2: 1.00, 1: 0.90}.get(level, 1.00)

def _build_tfidf(section_bodies: List[List[str]]) -> Tuple[List[Dict[str, float]], Dict[str, float]]:
    tf_list = [_tf(tokens) for tokens in section_bodies]
    df: Dict[str, int] = {}
    for tf in tf_list:
        for term in tf.keys():
            df[term] = df.get(term, 0) + 1
    N = float(len(tf_list)) or 1.0
    idf: Dict[str, float] = {}
    for term, cnt in df.items():
        idf[term] = math.log((N + 1.0) / (cnt + 1.0)) + 1.0
    return tf_list, idf

def _apply_idf(tf: Dict[str, float], idf: Dict[str, float]) -> Dict[str, float]:
    return {t: w * idf.get(t, 1.0) for t, w in tf.items()}

def _best_section_for_claim(claim_text: str, sections: List[Dict]) -> Tuple[str, str, int]:
    bodies_tokens = [_tokens(s.get("body", "")) for s in sections]
    tf_list, idf = _build_tfidf(bodies_tokens)
    claim_vec = _apply_idf(_tf(_tokens(claim_text)), idf)
    best_score = -1.0
    best = ("unknown", "unknown", "")

    for s, tf in zip(sections, tf_list):
        sec_vec = _apply_idf(tf, idf)
        sim = _cosine(claim_vec, sec_vec)
        sim *= _heading_penalty_factor(s.get("title", ""))
        sim *= _level_boost(int(s.get("level", 2)))
        if sim > best_score:
            best_score = sim
            best = (s.get("title", "unknown"), s.get("title", "unknown"), int(s.get("level", 2)) if s.get("level") else "")

    if best_score <= 0.0:
        claim_toks = set(_tokens(claim_text))
        best2 = ("unknown", "unknown", "")
        best2_score = -1
        for s in sections:
            overlap = len(claim_toks.intersection(set(_tokens(s.get("title", "")))))
            score = overlap * _heading_penalty_factor(s.get("title", "")) * _level_boost(int(s.get("level", 2)))
            if score > best2_score:
                best2_score = score
                best2 = (s.get("title", "unknown"), s.get("title", "unknown"), int(s.get("level", 2)) if s.get("level") else "")
        return best2

    return best

# ---------------------------
# PDF filename normalization & matching
# ---------------------------

_YEAR_RE = re.compile(r"\b(19|20)\d{2}[a-z]?\b", re.IGNORECASE)

def _norm_pdf_stem(name: str) -> str:
    s = (name or "")
    s = s.replace("-", "_")
    s = s.replace("&", " and ")
    s = s.translate(_PUNCT_TABLE_SP)
    s = re.sub(r"[^a-zA-Z0-9_ ]+", " ", s)
    s = s.lower()
    s = re.sub(r"\s+", " ", s).strip()
    s = s.replace(" ", "_")
    s = re.sub(r"_+", "_", s)
    return s

def _extract_year(s: str) -> Optional[str]:
    m = _YEAR_RE.search(s or "")
    return m.group(0).lower() if m else None

def _strip_et_al(s: str) -> str:
    return re.sub(r"\bet\.?\s*al\.?\b", "", s or "", flags=re.IGNORECASE)

def _primary_author_from_citation_author(cit_author: str) -> str:
    """
    Returns a conservative primary-author token from citation_author,
    prioritizing the first surname or org token.
    """
    a = _strip_et_al(cit_author or "")
    a = a.replace("&", " and ")
    a = re.sub(r"[\(\)\[\]\{\},.;:]", " ", a)
    a = re.sub(r"\s+", " ", a).strip()
    # Split "Masvaure & Fish" -> ["Masvaure", "Fish"], take first
    first_piece = re.split(r"\band\b|,|\u2013|\u2014|–|—", a, maxsplit=1, flags=re.IGNORECASE)[0].strip()
    base = first_piece.lower().replace("-", "_").replace("’", "").replace("'", "")
    base = re.sub(r"\s+", "_", base).strip("_")
    if not base:
        return ""
    # Reduce multi-token to last word too (surname), but keep full variant as primary
    return base

def _author_variants(author: str) -> List[str]:
    """
    Build robust author tokens.
    """
    a = _strip_et_al(author or "")
    a = a.replace("&", " and ")
    a = re.sub(r"[\(\)\[\]\{\},.;:]", " ", a)
    a = re.sub(r"\s+", " ", a).strip()
    first = re.split(r"\band\b|,|\u2013|\u2014|–|—", a, maxsplit=1, flags=re.IGNORECASE)[0].strip()
    tokens = first.split()
    surname = tokens[-1] if tokens else first
    base = first.lower()
    base = base.replace("-", "_").replace("’", "").replace("'", "")
    base = re.sub(r"\s+", "_", base).strip("_")
    variants = {base}
    if surname:
        s = surname.lower().replace("-", "_").replace("’", "").replace("'", "")
        s = re.sub(r"\s+", "_", s)
        variants.add(s)
    more = set()
    for v in list(variants):
        more.add(v.replace("_", ""))
    variants |= more
    prefixes = {"de", "van", "von", "da", "del", "di"}
    if "_" in base:
        parts = base.split("_")
        if parts[0] in prefixes and len(parts) > 1:
            variants.add("_".join(parts[1:]))
            variants.add("".join(parts[1:]))
    return [v for v in variants if v]

def _mine_authors_from_citation_text(citation_text: str) -> List[str]:
    if not citation_text:
        return []
    txt = _strip_et_al(citation_text)
    txt = txt.replace("&", " and ")
    txt = re.sub(r"[\(\)]", " ", txt)
    year = _extract_year(txt)
    if year:
        head = txt.split(year, 1)[0]
    else:
        head = txt
    chunks = re.split(r",|\band\b", head, flags=re.IGNORECASE)
    authors = []
    for c in chunks:
        c = _norm_pdf_stem(c)
        c = c.strip("_ ")
        if not c:
            continue
        parts = c.split("_")
        if len(parts) == 1:
            authors.append(parts[0])
        else:
            authors.append(parts[-1])
            authors.append(c)
    out = set()
    for a in authors:
        out.add(a)
        out.add(a.replace("_", ""))
    return [x for x in out if x]

def _split_multi_citation_parts(citation_text: str) -> List[str]:
    """
    Split "(A, 2022; B, 2018)" into ["A, 2022", "B, 2018"] robustly.
    """
    if not citation_text:
        return []
    t = citation_text.strip()
    t = t[1:-1] if t.startswith("(") and t.endswith(")") else t  # strip outer parens
    # now split on ';'
    parts = [p.strip() for p in t.split(";")]
    return [p for p in parts if p]

def _year_base(y: Optional[str]) -> Optional[str]:
    """Turn '2022a' -> '2022' for comparisons."""
    if not y:
        return None
    m = re.match(r"^((?:19|20)\d{2})", y.lower())
    return m.group(1) if m else y.lower()

def _build_pdf_index(pdf_dir: Path) -> List[Dict]:
    items: List[Dict] = []
    if not pdf_dir.exists():
        return items
    for p in pdf_dir.glob("*.pdf"):
        stem = p.stem
        stem_norm = _norm_pdf_stem(stem)
        items.append({
            "path": p,
            "stem_norm": stem_norm,
            "year": _extract_year(stem_norm),
        })
    return items

def _score_pdf_match(
    stem_norm: str,
    year: Optional[str],
    author_variants: List[str],
    want_year: Optional[str],
    primary_variant: Optional[str] = None
) -> int:
    """
    Stricter scoring:
      +10 exact year match (base year equality required)
      -999 if years differ (disqualify)
    """
    score = 0
    stem_year_base = _year_base(year)
    want_year_base = _year_base(want_year)

    if want_year_base and stem_year_base:
        if want_year_base == stem_year_base:
            score += 10
        else:
            return -999  # disqualify mismatched year entirely

    # author variant scoring stays the same
    if primary_variant:
        if re.search(rf"(?:^|_){re.escape(primary_variant)}(?:_|$)", stem_norm):
            score += 4
        elif primary_variant in stem_norm:
            score += 2

    for v in author_variants:
        if not v:
            continue
        if primary_variant and v == primary_variant:
            continue
        if re.search(rf"(?:^|_){re.escape(v)}(?:_|$)", stem_norm):
            score += 2
        elif v in stem_norm:
            score += 1

    return score

def _infer_year_for_target_author(citation_author: str, citation_text: str) -> Optional[str]:
    """
    In a multi-citation like "(Chirau et al., 2022; Kanyamuna et al., 2018)",
    if this row's citation_author ~ "Chirau et al.", return '2022';
    if ~ "Kanyamuna et al.", return '2018'.
    """
    if not citation_text:
        return None
    # get a primary token for matching inside parts
    primary = _primary_author_from_citation_author(citation_author)
    if not primary:
        return None
    primary_variants = {primary, primary.replace("_", "")}
    parts = _split_multi_citation_parts(citation_text)
    for part in parts:
        pn = _norm_pdf_stem(part)
        # does this part mention the primary author?
        if any(re.search(rf"(?:^|_){re.escape(v)}(?:_|$)", pn) for v in primary_variants):
            # return the year from this specific sub-citation
            y = _extract_year(pn)
            if y:
                return y
    # fallback: first year in the whole citation text
    return _extract_year(citation_text)

def resolve_pdf_path(row: Dict[str, str], pdf_index: List[Dict]) -> Optional[Path]:
    """
    Decide best matching PDF for a registry row. Return a Path or None.
    Priority & rules:
      - Prefer the FIRST author listed in citation_author (e.g., 'Masvaure' in 'Masvaure & Fish')
      - In multi-citations, pull the year that belongs to the target author segment
      - Penalize candidates with wrong year
    """
    cit_author = (row.get("citation_author") or "").strip()
    cit_text = (row.get("citation_text") or "").strip()
    cit_year = (str(row.get("citation_year") or "").strip() or None)

    # Primary author preference
    primary_variant = _primary_author_from_citation_author(cit_author)
    # Build author variants from citation_author + citation_text
    variants: List[str] = []
    if cit_author:
        variants += _author_variants(cit_author)
    if cit_text:
        variants += _mine_authors_from_citation_text(cit_text)

    # Add primary mentioned author as last resort
    pm_author = (row.get("primary_mentioned_author") or "").strip()
    pm_year = (str(row.get("primary_mentioned_year") or "").strip() or None)
    if pm_author:
        variants += _author_variants(pm_author)

    # Dedup variants, longer first
    seen = set()
    uniq_variants: List[str] = []
    for v in sorted(set(variants), key=lambda x: (-len(x), x)):
        if v not in seen:
            uniq_variants.append(v)
            seen.add(v)

    # Choose the desired year intelligently:
    # 1) If citation_year is present, start with it
    # 2) If multi-citation, try to infer the year that belongs to this row's target author
    # 3) Else, fallback to primary-mentioned year or first year in whole citation_text
    want_year = cit_year
    inferred = _infer_year_for_target_author(cit_author, cit_text)
    if inferred:
        want_year = inferred
    elif not want_year:
        want_year = pm_year or _extract_year(cit_text)

    # Score each PDF candidate
    best_score = -10**9
    best_item: Optional[Dict] = None
    for item in pdf_index:
        s = item["stem_norm"]
        y = item["year"]
        score = _score_pdf_match(
            stem_norm=s,
            year=y,
            author_variants=uniq_variants,
            want_year=want_year,
            primary_variant=primary_variant if primary_variant else None
        )
        if score > best_score:
            best_score = score
            best_item = item

    # Require a minimal signal
    if best_item and best_score >= 2:
        return best_item["path"]
    return None

# ---------------------------
# CSV IO
# ---------------------------

def _read_csv(path: Path) -> List[dict]:
    with open(path, "r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))

def _write_csv(path: Path, rows: List[dict], fieldnames: List[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames)
        w.writeheader()
        for r in rows:
            w.writerow(r)

# ---------------------------
# Public API
# ---------------------------

def enrich_registry_rows(raw_rows: List[dict], reviews_dir: Path, pdf_dir: Optional[Path] = None) -> List[dict]:
    """
    For each review_id in raw_rows:
      - open pilot_inputs/reviews/<review_id>.docx, parse H1..H4 and build bodies
      - match each claim_text to the best section via TF-IDF cosine
      - If source_pdf_path is blank or missing, resolve it against `pdf_dir`
    Append: section_canonical, section_level, research_question, section_title
    """
    doc_cache: Dict[str, Tuple[str, List[Dict]]] = {}
    pdf_index: List[Dict] = _build_pdf_index(pdf_dir) if pdf_dir else []

    enriched: List[dict] = []
    for r in raw_rows:
        rid = r.get("review_id", "")
        claim_text = r.get("claim_text", "") or ""

        if rid not in doc_cache:
            docx_path = reviews_dir / f"{rid}.docx"
            rq, sections = _load_docx_sections(docx_path)
            sections = [s for s in sections if s.get("level") in (2, 3, 4) or (s.get("level") == 1 and s.get("body"))]
            doc_cache[rid] = (rq, sections)

        rq, sections = doc_cache[rid]

        # Section mapping
        if sections:
            section_title, section_canonical, section_level = _best_section_for_claim(claim_text, sections)
        else:
            section_title, section_canonical, section_level = ("unknown", "unknown", "")

        out = dict(r)
        out["section_canonical"] = section_canonical
        out["section_level"] = section_level
        out["research_question"] = rq
        out["section_title"] = section_title

        # PDF resolution if missing/blank
        existing = (out.get("source_pdf_path") or "").strip()
        if (not existing) and pdf_index:
            m = resolve_pdf_path(out, pdf_index)
            if m:
                # store as relative path (from repo root)
                try:
                    rel = m.relative_to(Path.cwd())
                    out["source_pdf_path"] = str(rel).replace("\\", "/")
                except Exception:
                    out["source_pdf_path"] = str(m).replace("\\", "/")    
        else:
            # normalize already-populated paths too
            if existing:
                out["source_pdf_path"] = existing.replace("\\", "/")
        enriched.append(out)
    return enriched

def write_enriched_csv(enriched_rows: List[dict], out_path: Path) -> None:
    """
    Write CSV preserving original columns + 4 new enrichment columns at the end.
    """
    if not enriched_rows:
        _write_csv(out_path, [], [])
        return
    base_cols = list(enriched_rows[0].keys())
    for c in ["section_canonical", "section_level", "research_question", "section_title"]:
        if c in base_cols:
            base_cols.remove(c)
    fieldnames = base_cols + ["section_canonical", "section_level", "research_question", "section_title"]
    _write_csv(out_path, enriched_rows, fieldnames)