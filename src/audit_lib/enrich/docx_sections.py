from __future__ import annotations
from pathlib import Path
from typing import Tuple, List, Dict
import re


def norm_spaces(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip())


def load_docx_sections(docx_path: Path) -> Tuple[str, List[Dict]]:
    """Parse headings and bodies from a .docx review file.

    Returns:
        research_question (H1 title or ""),
        sections: list of dicts with keys ``level``, ``title``, ``body``.
    """
    try:
        from docx import Document  # type: ignore
    except ImportError as e:  # pragma: no cover - defensive
        raise RuntimeError("python-docx is required. pip install python-docx") from e

    if not docx_path.exists():
        return "", []

    doc = Document(str(docx_path))

    heads = []  # (idx, level, title)
    for i, p in enumerate(doc.paragraphs):
        style_name = (getattr(p.style, "name", "") or "").lower()
        if style_name.startswith("heading"):
            m = re.search(r"(\d+)", style_name)
            if not m:
                continue
            lvl = int(m.group(1))
            if 1 <= lvl <= 4:
                title = norm_spaces(p.text)
                if title:
                    heads.append((i, lvl, title))

    if not heads:
        return "", []

    rq = ""
    for _, lvl, title in heads:
        if lvl == 1:
            rq = title
            break

    sections: List[Dict] = []
    for j, (idx, lvl, title) in enumerate(heads):
        nxt = heads[j + 1][0] if j + 1 < len(heads) else len(doc.paragraphs)
        chunks = []
        for k in range(idx + 1, nxt):
            t = norm_spaces(doc.paragraphs[k].text)
            if t:
                chunks.append(t)
        body = " ".join(chunks)
        sections.append({"level": lvl, "title": title, "body": body})

    return rq, sections
